"""
resume.py — pull plain text out of an uploaded resume. In memory only.

Nothing here ever touches disk or the database. The extracted text lives in
Streamlit's session state (see view_profile() in app.py) for exactly as long
as the browser tab stays open, then it's gone — the founder's own words were
"make it so we're not tracking their personal information." The one honest
exception: the text IS sent to Anthropic to draft a reply, once, when the
person asks for one. That's disclosed in the privacy policy (legal.py), same
commit as this file.
"""
import io
import re

MAX_CHARS = 6000   # roughly two pages — enough signal, keeps the prompt cheap
# Belt and suspenders on top of .streamlit/config.toml's server-level
# maxUploadSize=10. That config caps what the browser can even SEND; this caps
# what this function will bother to PARSE, so a future caller that skips the
# config (a different entry point, a test harness) still can't hand pypdf a
# huge byte string on a small instance.
MAX_BYTES = 10 * 1024 * 1024


def extract_text(uploaded_file) -> str:
    """
    Best-effort plain text from a Streamlit UploadedFile.

    '' on anything unreadable OR too large (a scanned/image-only PDF, a
    corrupt file, something oversized) — callers should treat that as
    "nothing changed" and let drafting continue without it, never as an
    error that blocks the feature.
    """
    name = (uploaded_file.name or "").lower()
    try:
        data = uploaded_file.getvalue()
        if len(data) > MAX_BYTES:
            return ""
        if name.endswith(".pdf"):
            text = _from_pdf(data)
        elif name.endswith(".docx"):
            text = _from_docx(data)
        else:
            # .txt, and the fallback for anything else. NOT used for .docx: a
            # .docx is a zip archive, and decode(errors="ignore") on binary
            # doesn't raise — it silently returns garbage that LOOKS like text,
            # which would have gone straight into the AI prompt instead of
            # failing cleanly. Caught by testing against a real resume rather
            # than the synthetic PDFs this was first built and verified with.
            text = data.decode("utf-8", errors="ignore")
    except Exception:
        return ""
    text = re.sub(r"\n{3,}", "\n\n", text).strip()
    return text[:MAX_CHARS]


def _from_pdf(data: bytes) -> str:
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(data))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def _from_docx(data: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    # Resumes commonly put contact info, dates or skills in a table rather than
    # a paragraph (a two-column layout, a skills grid) — paragraphs alone would
    # silently drop that content.
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)
